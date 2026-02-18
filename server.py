#!/usr/bin/env python3
"""
MCP Document Server
Provides secure access to documents from a local directory for Claude AI
"""

import asyncio
import logging
import os
import time
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP
from mcp.types import Resource, TextContent, Tool

# Configure logging to stderr (never stdout for STDIO transport)
LOG_LEVEL = os.getenv('LOG_LEVEL', 'INFO').upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format='%(asctime)s - %(name)s - %(levelname)s - [%(funcName)s:%(lineno)d] %(message)s',
    handlers=[logging.StreamHandler()]  # Will use stderr by default
)
logger = logging.getLogger(__name__)
logger.info(f"Log level set to: {LOG_LEVEL}")

# Configuration from environment variables
DOCUMENTS_PATH = Path(os.getenv('DOCUMENTS_PATH', '/documents'))
MAX_FILE_SIZE = int(os.getenv('MAX_FILE_SIZE_MB', '10')) * 1024 * 1024  # Default 10MB
MAX_VIDEO_SIZE = int(os.getenv('MAX_VIDEO_SIZE_MB', '500')) * 1024 * 1024  # Default 500MB for videos
# Document extensions (text-extractable)
ALLOWED_EXTENSIONS = os.getenv('ALLOWED_EXTENSIONS', '.txt,.md,.pdf,.docx,.xlsx,.pptx,.csv,.json,.yaml,.yml,.log').split(',')

# Code/config extensions (readable as text)
CODE_EXTENSIONS = ['.py', '.js', '.sh', '.html', '.css', '.j2', '.proto', '.conf', '.vtt', '.xml', '.toml', '.ini', '.cfg']

# Image extensions
IMAGE_EXTENSIONS = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg', '.ico']

# Audio extensions
AUDIO_EXTENSIONS = ['.wav', '.mp3', '.m4a', '.ogg', '.flac', '.wma', '.aac', '.opus']

# Video extensions
VIDEO_EXTENSIONS = ['.mp4', '.mkv', '.avi', '.webm', '.mov', '.m4v']

# Archive/binary extensions (downloadable but not readable as text)
BINARY_EXTENSIONS = ['.zip', '.tgz', '.tar', '.gz', '.7z', '.rar', '.apk', '.bin', '.exe', '.msi', '.dmg', '.msg']

# All supported extensions for listing
ALL_EXTENSIONS = ALLOWED_EXTENSIONS + CODE_EXTENSIONS + IMAGE_EXTENSIONS + AUDIO_EXTENSIONS + VIDEO_EXTENSIONS + BINARY_EXTENSIONS
AUTH_TOKEN = os.getenv('MCP_AUTH_TOKEN', '')
PUBLIC_URL = os.getenv('MCP_PUBLIC_URL', '')  # e.g., https://frigate.taila099fd.ts.net
WHISPER_API_URL = os.getenv('WHISPER_API_URL', '')  # e.g., http://desktop:9000/v1/audio/transcriptions
WHISPER_SCAN_INTERVAL = int(os.getenv('WHISPER_SCAN_INTERVAL', '60'))  # seconds between scans
TRANSCRIPT_EXT = '.vtt'  # preferred transcript format

# Recordings subdirectory (audio scanner + speaker enrollment scoped here)
RECORDINGS_DIR = os.getenv('RECORDINGS_DIR', 'Recordings')

# Speaker identification settings
SPEAKERS_DIR = os.getenv('SPEAKERS_DIR', 'speakers')
SPEAKER_SIMILARITY_THRESHOLD = float(os.getenv('SPEAKER_SIMILARITY_THRESHOLD', '0.75'))
SPEAKER_MIN_AUDIO_SECONDS = 3.0  # minimum audio needed per speaker for embedding

# Track files currently being transcribed to avoid duplicate submissions
_transcription_in_progress: set[str] = set()
_transcription_errors: dict[str, str] = {}  # file_path -> last error message

# Speaker identification cache
_speaker_embeddings: dict[str, Any] = {}  # {name: {'embedding': np.array, 'mtime': float, 'paths': [...]}}
_resemblyzer_encoder = None  # lazy-loaded VoiceEncoder singleton

# Initialize FastMCP server with transport security settings
from mcp.server.transport_security import TransportSecuritySettings

# Allow connections from Tailscale Funnel and localhost
ALLOWED_HOSTS = os.getenv('MCP_ALLOWED_HOSTS', 'localhost,127.0.0.1').split(',')

transport_security = TransportSecuritySettings(
    allowed_hosts=ALLOWED_HOSTS,
    allowed_origins=['*'],  # Allow any origin for SSE connections
)

mcp = FastMCP("Document Server", transport_security=transport_security)

logger.info(f"Document server starting with path: {DOCUMENTS_PATH}")
logger.info(f"Allowed extensions: {ALLOWED_EXTENSIONS}")


def is_safe_path(base_path: Path, requested_path: Path) -> bool:
    """Ensure requested path is within the base documents directory"""
    try:
        resolved_base = base_path.resolve()
        resolved_requested = requested_path.resolve()
        return str(resolved_requested).startswith(str(resolved_base))
    except Exception:
        return False


def get_file_info(file_path: Path) -> dict[str, Any]:
    """Get file metadata"""
    stat = file_path.stat()
    suffix = file_path.suffix.lower()

    # Determine file type
    if suffix in [e.lower() for e in VIDEO_EXTENSIONS]:
        file_type = 'video'
    elif suffix in [e.lower() for e in AUDIO_EXTENSIONS]:
        file_type = 'audio'
    elif suffix in [e.lower() for e in IMAGE_EXTENSIONS]:
        file_type = 'image'
    elif suffix in [e.lower() for e in BINARY_EXTENSIONS]:
        file_type = 'binary'
    elif suffix in [e.lower() for e in CODE_EXTENSIONS]:
        file_type = 'code'
    else:
        file_type = 'document'

    return {
        'name': file_path.name,
        'path': str(file_path.relative_to(DOCUMENTS_PATH)),
        'size': stat.st_size,
        'modified': stat.st_mtime,
        'extension': file_path.suffix,
        'type': file_type
    }


def get_video_metadata(file_path: Path) -> dict[str, Any]:
    """Extract video metadata using ffprobe"""
    import subprocess
    import json as json_module

    try:
        cmd = [
            'ffprobe', '-v', 'quiet',
            '-print_format', 'json',
            '-show_format', '-show_streams',
            str(file_path)
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if result.returncode != 0:
            logger.warning(f"ffprobe failed for {file_path}: {result.stderr}")
            return {'error': 'Could not extract video metadata'}

        data = json_module.loads(result.stdout)

        # Extract useful info
        format_info = data.get('format', {})
        streams = data.get('streams', [])

        video_stream = next((s for s in streams if s.get('codec_type') == 'video'), None)
        audio_stream = next((s for s in streams if s.get('codec_type') == 'audio'), None)

        metadata = {
            'duration': float(format_info.get('duration', 0)),
            'duration_formatted': format_duration(float(format_info.get('duration', 0))),
            'size_bytes': int(format_info.get('size', 0)),
            'bit_rate': int(format_info.get('bit_rate', 0)) if format_info.get('bit_rate') else None,
            'format_name': format_info.get('format_name'),
        }

        if video_stream:
            metadata['video'] = {
                'codec': video_stream.get('codec_name'),
                'width': video_stream.get('width'),
                'height': video_stream.get('height'),
                'fps': eval(video_stream.get('r_frame_rate', '0/1')) if video_stream.get('r_frame_rate') else None,
            }

        if audio_stream:
            metadata['audio'] = {
                'codec': audio_stream.get('codec_name'),
                'channels': audio_stream.get('channels'),
                'sample_rate': audio_stream.get('sample_rate'),
            }

        # Include any embedded metadata tags
        tags = format_info.get('tags', {})
        if tags:
            metadata['tags'] = {k: v for k, v in tags.items() if k.lower() in ['title', 'artist', 'date', 'creation_time', 'comment']}

        return metadata

    except subprocess.TimeoutExpired:
        logger.error(f"ffprobe timed out for {file_path}")
        return {'error': 'Video metadata extraction timed out'}
    except Exception as e:
        logger.error(f"Error getting video metadata: {e}", exc_info=True)
        return {'error': str(e)}


def format_duration(seconds: float) -> str:
    """Format duration in seconds to HH:MM:SS"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    return f"{minutes}:{secs:02d}"


def _get_voice_encoder():
    """Lazy-load the resemblyzer VoiceEncoder on CPU. Returns None if unavailable."""
    global _resemblyzer_encoder
    if _resemblyzer_encoder is not None:
        return _resemblyzer_encoder
    try:
        from resemblyzer import VoiceEncoder
        logger.info("Loading resemblyzer VoiceEncoder (CPU)...")
        _resemblyzer_encoder = VoiceEncoder("cpu")
        logger.info("VoiceEncoder loaded successfully")
        return _resemblyzer_encoder
    except ImportError:
        logger.warning("resemblyzer not installed - speaker identification disabled")
        return None
    except Exception as e:
        logger.error(f"Failed to load VoiceEncoder: {e}", exc_info=True)
        return None


def _get_speakers_path() -> Path:
    """Return the speakers enrollment directory path (inside recordings dir)."""
    return DOCUMENTS_PATH / RECORDINGS_DIR / SPEAKERS_DIR


def _scan_speaker_profiles() -> dict[str, list[Path]]:
    """Scan speakers/ directory, return {name: [audio_paths]}.

    Supports flat layout (speakers/Sarah.wav) and subdirectory layout
    (speakers/Sarah/clip1.wav).
    """
    speakers_path = _get_speakers_path()
    if not speakers_path.exists() or not speakers_path.is_dir():
        return {}

    audio_exts = set(e.lower() for e in AUDIO_EXTENSIONS)
    profiles: dict[str, list[Path]] = {}

    for entry in speakers_path.iterdir():
        if entry.is_file() and entry.suffix.lower() in audio_exts:
            # Flat layout: filename stem is the speaker name
            name = entry.stem
            profiles.setdefault(name, []).append(entry)
        elif entry.is_dir():
            # Subdirectory layout: directory name is the speaker name
            name = entry.name
            clips = [f for f in entry.iterdir() if f.is_file() and f.suffix.lower() in audio_exts]
            if clips:
                profiles.setdefault(name, []).extend(clips)

    return profiles


def _load_speaker_embeddings() -> dict[str, Any]:
    """Load/cache speaker embeddings. Recomputes when enrollment files change (mtime check)."""
    global _speaker_embeddings
    import numpy as np

    profiles = _scan_speaker_profiles()
    if not profiles:
        _speaker_embeddings = {}
        return _speaker_embeddings

    encoder = _get_voice_encoder()
    if encoder is None:
        return {}

    from resemblyzer import preprocess_wav

    updated = False
    current_names = set(profiles.keys())
    cached_names = set(_speaker_embeddings.keys())

    # Remove profiles that no longer exist
    for gone in cached_names - current_names:
        del _speaker_embeddings[gone]
        updated = True

    for name, paths in profiles.items():
        # Check if any file has changed
        max_mtime = max(p.stat().st_mtime for p in paths)
        cached = _speaker_embeddings.get(name)
        if cached and cached['mtime'] >= max_mtime and cached['paths'] == [str(p) for p in paths]:
            continue  # Cache is still valid

        # Recompute embedding for this speaker
        try:
            all_wavs = []
            for clip_path in paths:
                wav = preprocess_wav(clip_path)
                if len(wav) / 16000 >= SPEAKER_MIN_AUDIO_SECONDS:
                    all_wavs.append(wav)
                else:
                    logger.warning(f"Speaker clip too short (<{SPEAKER_MIN_AUDIO_SECONDS}s), skipping: {clip_path}")

            if not all_wavs:
                logger.warning(f"No valid clips for speaker '{name}', skipping")
                continue

            # Compute embedding from all clips concatenated
            combined = np.concatenate(all_wavs)
            embedding = encoder.embed_utterance(combined)
            _speaker_embeddings[name] = {
                'embedding': embedding,
                'mtime': max_mtime,
                'paths': [str(p) for p in paths],
            }
            updated = True
            logger.info(f"Computed speaker embedding for '{name}' from {len(all_wavs)} clip(s)")
        except Exception as e:
            logger.error(f"Error computing embedding for speaker '{name}': {e}", exc_info=True)

    if updated:
        logger.info(f"Speaker embeddings updated: {len(_speaker_embeddings)} profile(s)")

    return _speaker_embeddings


def _extract_speaker_audio_segment(audio_path: Path, start: float, end: float) -> 'np.ndarray | None':
    """Use ffmpeg to extract an audio segment as a 16kHz mono numpy array."""
    import subprocess
    import tempfile
    import numpy as np

    duration = end - start
    if duration < 0.5:
        return None

    try:
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=True) as tmp:
            tmp_path = tmp.name

        cmd = [
            'ffmpeg', '-hide_banner', '-loglevel', 'warning',
            '-i', str(audio_path),
            '-ss', str(start),
            '-t', str(duration),
            '-ar', '16000',
            '-ac', '1',
            '-f', 'wav',
            '-y', tmp_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            logger.warning(f"ffmpeg segment extraction failed: {result.stderr}")
            return None

        import soundfile as sf
        wav_data, sr = sf.read(tmp_path, dtype='float32')
        os.unlink(tmp_path)
        return wav_data
    except Exception as e:
        logger.warning(f"Error extracting audio segment [{start:.1f}-{end:.1f}]: {e}")
        # Clean up temp file if it exists
        try:
            os.unlink(tmp_path)
        except (OSError, UnboundLocalError):
            pass
        return None


def _identify_speakers_in_segments(
    audio_path: Path,
    segments: list[dict],
    speaker_embeds: dict[str, Any],
) -> dict[str, str]:
    """Match anonymous speaker IDs (SPEAKER_00 etc.) to enrolled profiles.

    Returns a mapping like {'SPEAKER_00': 'Sarah', 'SPEAKER_01': 'SPEAKER_01'}.
    """
    import numpy as np

    encoder = _get_voice_encoder()
    if encoder is None:
        return {}

    # Collect segments per anonymous speaker
    speaker_segments: dict[str, list[dict]] = {}
    for seg in segments:
        spk = seg.get('speaker')
        if spk:
            speaker_segments.setdefault(spk, []).append(seg)

    speaker_map: dict[str, str] = {}
    enrolled_names = list(speaker_embeds.keys())
    enrolled_vecs = np.array([speaker_embeds[n]['embedding'] for n in enrolled_names])

    for anon_id, segs in speaker_segments.items():
        # Collect enough audio for this anonymous speaker (target ~10s, at least 3s)
        collected_wavs = []
        total_duration = 0.0

        # Sort by duration descending so we pick the longest segments first
        segs_sorted = sorted(segs, key=lambda s: s.get('end', 0) - s.get('start', 0), reverse=True)
        for seg in segs_sorted:
            start = seg.get('start', 0)
            end = seg.get('end', 0)
            wav = _extract_speaker_audio_segment(audio_path, start, end)
            if wav is not None and len(wav) > 0:
                collected_wavs.append(wav)
                total_duration += len(wav) / 16000
                if total_duration >= 10.0:
                    break

        if total_duration < SPEAKER_MIN_AUDIO_SECONDS:
            logger.debug(f"Not enough audio for {anon_id} ({total_duration:.1f}s), keeping anonymous")
            speaker_map[anon_id] = anon_id
            continue

        try:
            combined = np.concatenate(collected_wavs)
            anon_embedding = encoder.embed_utterance(combined)

            # Cosine similarity against all enrolled profiles
            similarities = np.dot(enrolled_vecs, anon_embedding) / (
                np.linalg.norm(enrolled_vecs, axis=1) * np.linalg.norm(anon_embedding)
            )

            best_idx = int(np.argmax(similarities))
            best_score = float(similarities[best_idx])
            best_name = enrolled_names[best_idx]

            logger.debug(f"{anon_id}: best match = '{best_name}' (score={best_score:.3f}, threshold={SPEAKER_SIMILARITY_THRESHOLD})")

            if best_score >= SPEAKER_SIMILARITY_THRESHOLD:
                speaker_map[anon_id] = best_name
                logger.info(f"Identified {anon_id} as '{best_name}' (similarity={best_score:.3f})")
            else:
                speaker_map[anon_id] = anon_id
                logger.info(f"No match for {anon_id} (best='{best_name}' at {best_score:.3f})")
        except Exception as e:
            logger.error(f"Error identifying {anon_id}: {e}", exc_info=True)
            speaker_map[anon_id] = anon_id

    return speaker_map


def _segments_to_labeled_vtt(segments: list[dict], speaker_map: dict[str, str]) -> str:
    """Convert WhisperX JSON segments + speaker map to a WebVTT string with real names."""
    lines = ["WEBVTT", ""]

    for seg in segments:
        start = seg.get('start', 0)
        end = seg.get('end', 0)
        text = seg.get('text', '').strip()
        speaker = seg.get('speaker', '')

        if not text:
            continue

        # Format timestamps as HH:MM:SS.mmm
        def fmt_ts(seconds: float) -> str:
            h = int(seconds // 3600)
            m = int((seconds % 3600) // 60)
            s = seconds % 60
            return f"{h:02d}:{m:02d}:{s:06.3f}"

        # Map anonymous speaker to real name
        label = speaker_map.get(speaker, speaker)
        labeled_text = f"[{label}]: {text}" if label else text

        lines.append(f"{fmt_ts(start)} --> {fmt_ts(end)}")
        lines.append(labeled_text)
        lines.append("")

    return "\n".join(lines)


def _build_whisper_url(original_url: str, target_output: str) -> str:
    """Swap the output format parameter in a Whisper API URL.

    e.g. change output=vtt to output=json (or vice versa).
    """
    import urllib.parse

    parsed = urllib.parse.urlparse(original_url)
    params = urllib.parse.parse_qs(parsed.query)
    params['output'] = [target_output]
    new_query = urllib.parse.urlencode(params, doseq=True)
    return urllib.parse.urlunparse(parsed._replace(query=new_query))


async def _transcribe_file(audio_path: Path) -> bool:
    """Internal: send a single audio file to the Whisper API and save the transcript.

    When speaker enrollment profiles exist, requests JSON output from Whisper,
    identifies speakers via voice embeddings, and generates labeled VTT.
    Otherwise falls back to standard VTT output.

    Returns True on success, False on failure."""
    import aiohttp
    import json as json_module

    rel_path = str(audio_path.relative_to(DOCUMENTS_PATH))

    if rel_path in _transcription_in_progress:
        logger.debug(f"Skipping {rel_path} - transcription already in progress")
        return False

    _transcription_in_progress.add(rel_path)
    start_time = time.time()
    logger.info(f"Background transcription starting: {rel_path} ({audio_path.stat().st_size / 1024 / 1024:.1f} MB)")

    # Determine if we should do speaker identification
    speaker_embeds = _load_speaker_embeddings()
    use_speaker_id = bool(speaker_embeds)
    if use_speaker_id:
        whisper_url = _build_whisper_url(WHISPER_API_URL, 'json')
        logger.info(f"Speaker profiles found ({len(speaker_embeds)}), requesting JSON output for speaker ID")
    else:
        whisper_url = WHISPER_API_URL
        logger.debug("No speaker profiles, using standard transcription output")

    try:
        file_handle = open(audio_path, 'rb')
        try:
            async with aiohttp.ClientSession() as session:
                data = aiohttp.FormData()
                data.add_field('audio_file', file_handle,
                              filename=audio_path.name,
                              content_type='audio/mpeg')

                async with session.post(whisper_url, data=data,
                                       timeout=aiohttp.ClientTimeout(total=14400)) as resp:
                    if resp.status != 200:
                        error_text = await resp.text()
                        logger.error(f"Whisper API returned {resp.status} for {rel_path}: {error_text}")
                        _transcription_errors[rel_path] = f"API error {resp.status}: {error_text[:200]}"
                        return False

                    response_text = await resp.text()
        finally:
            file_handle.close()

        # Process the response based on whether we're doing speaker ID
        if use_speaker_id:
            try:
                whisper_json = json_module.loads(response_text)
                segments = whisper_json.get('segments', [])

                if segments:
                    speaker_map = _identify_speakers_in_segments(audio_path, segments, speaker_embeds)
                    transcript_text = _segments_to_labeled_vtt(segments, speaker_map)
                    labeled_count = sum(1 for v in speaker_map.values() if not v.startswith('SPEAKER_'))
                    logger.info(f"Speaker ID complete: {labeled_count}/{len(speaker_map)} speakers identified")
                else:
                    logger.warning(f"No segments in Whisper JSON for {rel_path}, saving raw response")
                    transcript_text = response_text
            except (json_module.JSONDecodeError, KeyError) as e:
                logger.warning(f"Failed to parse Whisper JSON for speaker ID ({e}), saving raw response")
                transcript_text = response_text
        else:
            transcript_text = response_text

        # Save transcript alongside original file
        transcript_path = audio_path.with_suffix(TRANSCRIPT_EXT)
        with open(transcript_path, 'w', encoding='utf-8') as f:
            f.write(transcript_text)

        elapsed = time.time() - start_time
        transcript_size = len(transcript_text.encode('utf-8'))
        logger.info(f"Background transcription completed: {rel_path} -> {transcript_path.name} "
                    f"({transcript_size} bytes) in {elapsed:.1f}s")

        # Clear any previous error
        _transcription_errors.pop(rel_path, None)
        return True

    except asyncio.TimeoutError:
        elapsed = time.time() - start_time
        logger.error(f"Background transcription timed out for {rel_path} after {elapsed:.1f}s")
        _transcription_errors[rel_path] = f"Timed out after {elapsed:.0f}s"
        return False
    except aiohttp.ClientError as e:
        logger.error(f"Background transcription connection error for {rel_path}: {e}")
        _transcription_errors[rel_path] = f"Connection error: {e}"
        return False
    except Exception as e:
        logger.error(f"Background transcription error for {rel_path}: {e}", exc_info=True)
        _transcription_errors[rel_path] = str(e)
        return False
    finally:
        _transcription_in_progress.discard(rel_path)


async def _audio_scanner_loop():
    """Background loop that scans for audio files without transcripts and transcribes them."""
    logger.info(f"Audio scanner started (interval: {WHISPER_SCAN_INTERVAL}s)")

    # Wait a bit on startup before first scan
    await asyncio.sleep(5)

    while True:
        try:
            if not WHISPER_API_URL:
                logger.debug("Audio scanner: Whisper API not configured, sleeping")
                await asyncio.sleep(WHISPER_SCAN_INTERVAL)
                continue

            audio_exts_lower = [e.lower() for e in AUDIO_EXTENSIONS]
            speakers_path = _get_speakers_path()
            recordings_path = DOCUMENTS_PATH / RECORDINGS_DIR
            untranscribed = []

            if not recordings_path.exists():
                logger.debug(f"Audio scanner: recordings dir not found ({RECORDINGS_DIR}), sleeping")
                await asyncio.sleep(WHISPER_SCAN_INTERVAL)
                continue

            for file_path in recordings_path.rglob('*'):
                if not file_path.is_file():
                    continue
                if file_path.suffix.lower() not in audio_exts_lower:
                    continue

                # Skip audio files inside the speakers enrollment directory
                try:
                    file_path.relative_to(speakers_path)
                    continue  # Inside speakers/ dir, skip it
                except ValueError:
                    pass  # Not inside speakers/ dir, proceed

                rel_path = str(file_path.relative_to(DOCUMENTS_PATH))

                # Skip if already in progress
                if rel_path in _transcription_in_progress:
                    continue

                # Check if transcript exists
                transcript_path = file_path.with_suffix(TRANSCRIPT_EXT)
                if not transcript_path.exists():
                    untranscribed.append(file_path)

            if untranscribed:
                logger.info(f"Audio scanner: found {len(untranscribed)} untranscribed audio file(s)")
                for audio_path in untranscribed:
                    # Transcribe one at a time to avoid overloading the GPU
                    await _transcribe_file(audio_path)
            else:
                logger.debug("Audio scanner: no untranscribed audio files found")

        except Exception as e:
            logger.error(f"Audio scanner error: {e}", exc_info=True)

        await asyncio.sleep(WHISPER_SCAN_INTERVAL)


@mcp.tool()
async def list_documents(
    subdirectory: str = "",
    recursive: bool = False
) -> str:
    """
    List available documents in the document directory.

    Args:
        subdirectory: Optional subdirectory to list (relative path)
        recursive: If True, list files recursively

    Returns:
        JSON string with list of documents and their metadata
    """
    start_time = time.time()
    logger.info(f"list_documents called: subdirectory='{subdirectory}', recursive={recursive}")

    try:
        search_path = DOCUMENTS_PATH / subdirectory if subdirectory else DOCUMENTS_PATH
        logger.debug(f"Search path resolved to: {search_path}")

        if not is_safe_path(DOCUMENTS_PATH, search_path):
            logger.warning(f"Path safety check failed for: {search_path}")
            return f"Error: Access denied - path outside allowed directory"

        if not search_path.exists():
            logger.warning(f"Directory not found: {search_path}")
            return f"Error: Directory not found: {subdirectory}"

        documents = []

        if recursive:
            pattern = '**/*'
        else:
            pattern = '*'

        logger.debug(f"Scanning with pattern: {pattern}")
        for file_path in search_path.glob(pattern):
            if file_path.is_file() and (file_path.suffix.lower() in [e.lower() for e in ALL_EXTENSIONS]):
                documents.append(get_file_info(file_path))
                logger.debug(f"Found document: {file_path.name}")

        # Sort by modification time (newest first)
        documents.sort(key=lambda x: x['modified'], reverse=True)

        elapsed = time.time() - start_time
        logger.info(f"list_documents completed: found {len(documents)} documents in {elapsed:.2f}s")

        import json
        return json.dumps({
            'directory': str(subdirectory or '/'),
            'total_files': len(documents),
            'documents': documents
        }, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error listing documents after {elapsed:.2f}s: {e}", exc_info=True)
        return f"Error listing documents: {str(e)}"


@mcp.tool()
async def read_document(file_path: str, max_chars: int = 500000) -> str:
    """
    Read the contents of a document.

    Args:
        file_path: Relative path to the document
        max_chars: Maximum characters to return (default 500000)

    Returns:
        Document contents as string
    """
    start_time = time.time()
    logger.info(f"read_document called: file_path='{file_path}', max_chars={max_chars}")

    try:
        full_path = DOCUMENTS_PATH / file_path
        logger.debug(f"Full path resolved to: {full_path}")

        if not is_safe_path(DOCUMENTS_PATH, full_path):
            logger.warning(f"Path safety check failed for: {full_path}")
            return f"Error: Access denied - path outside allowed directory"

        if not full_path.exists():
            logger.warning(f"File not found: {full_path}")
            return f"Error: File not found: {file_path}"

        suffix = full_path.suffix.lower()
        is_video = suffix in [e.lower() for e in VIDEO_EXTENSIONS]
        is_audio = suffix in [e.lower() for e in AUDIO_EXTENSIONS]
        is_image = suffix in [e.lower() for e in IMAGE_EXTENSIONS]
        is_binary = suffix in [e.lower() for e in BINARY_EXTENSIONS]
        is_code = suffix in [e.lower() for e in CODE_EXTENSIONS]
        is_document = suffix in [e.lower() for e in ALLOWED_EXTENSIONS]

        if not (is_video or is_audio or is_image or is_binary or is_code or is_document):
            logger.warning(f"Extension not allowed: {full_path.suffix}")
            return f"Error: File type not allowed: {full_path.suffix}"

        # Check file size (videos/audio/images/binaries have higher limit)
        file_size = full_path.stat().st_size
        logger.debug(f"File size: {file_size} bytes")
        size_limit = MAX_VIDEO_SIZE if (is_video or is_audio or is_image or is_binary) else MAX_FILE_SIZE
        if file_size > size_limit:
            logger.warning(f"File too large: {file_size} > {size_limit}")
            return f"Error: File too large ({file_size} bytes, max {size_limit})"

        # Read file based on extension
        if full_path.suffix == '.pdf':
            logger.debug("Reading PDF file")
            try:
                import pypdf
                reader = pypdf.PdfReader(full_path)
                logger.debug(f"PDF has {len(reader.pages)} pages")
                text = ""
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    logger.debug(f"Page {i+1}: extracted {len(page_text)} chars")
                content = text[:max_chars]
            except ImportError as e:
                logger.error(f"pypdf import failed: {e}")
                return "Error: PDF support not installed. Install pypdf package."
            except Exception as e:
                logger.error(f"PDF parsing error: {e}", exc_info=True)
                return f"Error reading PDF: {str(e)}"

        elif full_path.suffix in ['.docx']:
            logger.debug("Reading DOCX file")
            try:
                import docx
                doc = docx.Document(full_path)
                logger.debug(f"DOCX has {len(doc.paragraphs)} paragraphs")
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                content = text[:max_chars]
            except ImportError as e:
                logger.error(f"python-docx import failed: {e}")
                return "Error: DOCX support not installed. Install python-docx package."
            except Exception as e:
                logger.error(f"DOCX parsing error: {e}", exc_info=True)
                return f"Error reading DOCX: {str(e)}"

        elif full_path.suffix in ['.xlsx']:
            logger.debug("Reading XLSX file")
            try:
                import openpyxl
                wb = openpyxl.load_workbook(full_path)
                logger.debug(f"XLSX has {len(wb.sheetnames)} sheets: {wb.sheetnames}")
                text = f"Excel file with {len(wb.sheetnames)} sheets: {', '.join(wb.sheetnames)}\n\n"
                for sheet_name in wb.sheetnames[:5]:  # First 5 sheets
                    ws = wb[sheet_name]
                    row_count = 0
                    text += f"\n=== Sheet: {sheet_name} ===\n"
                    for row in ws.iter_rows():  # All rows (character limit still applies)
                        text += "\t".join([str(cell.value) if cell.value is not None else "" for cell in row]) + "\n"
                        row_count += 1
                    logger.debug(f"Sheet '{sheet_name}': read {row_count} rows")
                content = text[:max_chars]
            except ImportError as e:
                logger.error(f"openpyxl import failed: {e}")
                return "Error: XLSX support not installed. Install openpyxl package."
            except Exception as e:
                logger.error(f"XLSX parsing error: {e}", exc_info=True)
                return f"Error reading XLSX: {str(e)}"

        elif is_video:
            # Video files - return metadata
            logger.debug("Reading video file metadata")
            metadata = get_video_metadata(full_path)

            elapsed = time.time() - start_time
            logger.info(f"read_document (video) completed: got metadata, took {elapsed:.2f}s")

            import json
            result = {
                'file_path': file_path,
                'size': file_size,
                'extension': full_path.suffix,
                'type': 'video',
                'metadata': metadata,
                'content': f"Video file: {full_path.name}\n"
                          f"Duration: {metadata.get('duration_formatted', 'unknown')}\n"
                          f"Size: {file_size / 1024 / 1024:.1f} MB\n"
                          f"Format: {metadata.get('format_name', 'unknown')}\n"
            }

            if 'video' in metadata:
                v = metadata['video']
                result['content'] += f"Resolution: {v.get('width')}x{v.get('height')}\n"
                result['content'] += f"Video codec: {v.get('codec')}\n"
                if v.get('fps'):
                    result['content'] += f"FPS: {v.get('fps'):.2f}\n"

            if 'audio' in metadata:
                a = metadata['audio']
                result['content'] += f"Audio codec: {a.get('codec')}\n"
                result['content'] += f"Audio channels: {a.get('channels')}\n"

            if metadata.get('tags'):
                result['content'] += f"\nTags: {metadata['tags']}\n"

            result['content'] += "\nUse extract_video_text() to extract text content via OCR."
            result['content'] += "\nUse download_file() to get a download URL."
            result['content'] += "\nUse download_file_chunk() to download via base64 chunks."

            return json.dumps(result, indent=2)

        elif is_audio:
            # Audio files - return metadata and transcription status
            logger.debug("Reading audio file metadata")
            import json

            rel_path = str(full_path.relative_to(DOCUMENTS_PATH))

            # Check if a transcript already exists
            transcript_path = full_path.with_suffix(TRANSCRIPT_EXT)
            has_transcript = transcript_path.exists()
            is_in_progress = rel_path in _transcription_in_progress
            last_error = _transcription_errors.get(rel_path)

            elapsed = time.time() - start_time
            logger.info(f"read_document (audio) completed: got metadata, took {elapsed:.2f}s")

            result = {
                'file_path': file_path,
                'size': file_size,
                'extension': full_path.suffix,
                'type': 'audio',
                'content': f"Audio file: {full_path.name}\n"
                          f"Size: {file_size / 1024 / 1024:.1f} MB\n"
                          f"Format: {full_path.suffix.upper()}\n"
            }

            if has_transcript:
                result['content'] += f"\nTranscript available: {transcript_path.name}\n"
                result['content'] += "Use read_document() to read the transcript, or use transcribe_audio() to retrieve it.\n"
                result['transcript_path'] = str(transcript_path.relative_to(DOCUMENTS_PATH))
            elif is_in_progress:
                result['content'] += "\nTranscription is currently in progress. Check back shortly.\n"
            else:
                result['content'] += "\nNo transcript yet. The background scanner will transcribe this file automatically.\n"
                if last_error:
                    result['content'] += f"Last transcription error: {last_error}\n"
                if not WHISPER_API_URL:
                    result['content'] += "Warning: Whisper API not configured (WHISPER_API_URL not set).\n"

            result['content'] += "\nUse download_file() to get a download URL.\n"
            result['content'] += "Use download_file_chunk() to download via base64 chunks."

            return json.dumps(result, indent=2)

        elif is_image:
            # Image files - return metadata and download instructions
            logger.debug("Reading image file metadata")
            import json

            elapsed = time.time() - start_time
            logger.info(f"read_document (image) completed: got metadata, took {elapsed:.2f}s")

            result = {
                'file_path': file_path,
                'size': file_size,
                'extension': full_path.suffix,
                'type': 'image',
                'content': f"Image file: {full_path.name}\n"
                          f"Size: {file_size / 1024:.1f} KB\n"
                          f"Format: {full_path.suffix.upper()}\n\n"
                          f"Use download_file() to get a download URL.\n"
                          f"Use download_file_chunk() to download via base64 chunks."
            }

            return json.dumps(result, indent=2)

        elif is_binary:
            # Binary/archive files - return metadata and download instructions
            logger.debug("Reading binary file metadata")
            import json

            elapsed = time.time() - start_time
            logger.info(f"read_document (binary) completed: got metadata, took {elapsed:.2f}s")

            result = {
                'file_path': file_path,
                'size': file_size,
                'extension': full_path.suffix,
                'type': 'binary',
                'content': f"Binary file: {full_path.name}\n"
                          f"Size: {file_size / 1024 / 1024:.1f} MB\n"
                          f"Format: {full_path.suffix.upper()}\n\n"
                          f"This file cannot be read as text.\n"
                          f"Use download_file() to get a download URL.\n"
                          f"Use download_file_chunk() to download via base64 chunks."
            }

            return json.dumps(result, indent=2)

        elif is_code:
            # Code files - read as text
            logger.debug(f"Reading code file with extension: {full_path.suffix}")
            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(max_chars)

        else:
            # Text-based document files
            logger.debug(f"Reading text file with extension: {full_path.suffix}")
            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(max_chars)

        truncated = len(content) >= max_chars
        elapsed = time.time() - start_time
        logger.info(f"read_document completed: {len(content)} chars, truncated={truncated}, took {elapsed:.2f}s")

        result = {
            'file_path': file_path,
            'size': file_size,
            'extension': full_path.suffix,
            'content': content,
            'truncated': truncated
        }

        import json
        return json.dumps(result, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error reading document {file_path} after {elapsed:.2f}s: {e}", exc_info=True)
        return f"Error reading document: {str(e)}"


@mcp.tool()
async def search_documents(
    query: str,
    file_extension: str = "",
    case_sensitive: bool = False
) -> str:
    """
    Search for documents containing specific text.

    Args:
        query: Text to search for
        file_extension: Optional file extension filter (e.g., '.txt')
        case_sensitive: Whether search should be case-sensitive

    Returns:
        JSON string with matching documents and snippets
    """
    start_time = time.time()
    logger.info(f"search_documents called: query='{query[:50]}...', extension='{file_extension}', case_sensitive={case_sensitive}")

    try:
        if not query:
            logger.warning("Empty search query")
            return "Error: Search query cannot be empty"

        search_query = query if case_sensitive else query.lower()
        matches = []
        files_scanned = 0
        files_skipped = 0

        for file_path in DOCUMENTS_PATH.rglob('*'):
            if not file_path.is_file():
                continue

            if file_path.suffix not in ALLOWED_EXTENSIONS:
                files_skipped += 1
                continue

            if file_extension and file_path.suffix != file_extension:
                files_skipped += 1
                continue

            try:
                # Only search text-based files for now
                if file_path.suffix in ['.txt', '.md', '.json', '.yaml', '.yml', '.log', '.csv']:
                    files_scanned += 1
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        search_content = content if case_sensitive else content.lower()

                        if search_query in search_content:
                            # Find snippet around match
                            index = search_content.find(search_query)
                            start = max(0, index - 100)
                            end = min(len(content), index + len(query) + 100)
                            snippet = content[start:end]

                            matches.append({
                                'path': str(file_path.relative_to(DOCUMENTS_PATH)),
                                'size': file_path.stat().st_size,
                                'snippet': snippet.strip()
                            })
                            logger.debug(f"Match found in: {file_path.name}")
                else:
                    files_skipped += 1
            except Exception as e:
                logger.warning(f"Error searching {file_path}: {e}")
                continue

        elapsed = time.time() - start_time
        logger.info(f"search_documents completed: {len(matches)} matches in {files_scanned} files ({files_skipped} skipped), took {elapsed:.2f}s")

        import json
        return json.dumps({
            'query': query,
            'total_matches': len(matches),
            'matches': matches[:50]  # Limit to 50 results
        }, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error searching documents after {elapsed:.2f}s: {e}", exc_info=True)
        return f"Error searching documents: {str(e)}"


@mcp.tool()
async def extract_video_text(
    file_path: str,
    fps: float = 0.5,
    similarity_threshold: int = 50
) -> str:
    """
    Extract text from a video file using OCR. Useful for videos of scrolling content
    like Teams chats, emails, or documents.

    Args:
        file_path: Relative path to the video file
        fps: Frames per second to extract (default 0.5 = 1 frame every 2 seconds)
        similarity_threshold: Percent line overlap to consider duplicate (default 50)

    Returns:
        JSON string with extracted text and metadata
    """
    import subprocess
    import tempfile
    import shutil
    from difflib import SequenceMatcher

    start_time = time.time()
    logger.info(f"extract_video_text called: file_path='{file_path}', fps={fps}, threshold={similarity_threshold}")

    try:
        import json

        full_path = DOCUMENTS_PATH / file_path
        logger.debug(f"Full path resolved to: {full_path}")

        if not is_safe_path(DOCUMENTS_PATH, full_path):
            logger.warning(f"Path safety check failed for: {full_path}")
            return json.dumps({'success': False, 'error': 'Access denied - path outside allowed directory'})

        if not full_path.exists():
            logger.warning(f"File not found: {full_path}")
            return json.dumps({'success': False, 'error': f'File not found: {file_path}'})

        if full_path.suffix.lower() not in VIDEO_EXTENSIONS:
            return json.dumps({'success': False, 'error': f'Not a video file: {full_path.suffix}'})

        # Check file size
        file_size = full_path.stat().st_size
        if file_size > MAX_VIDEO_SIZE:
            return json.dumps({'success': False, 'error': f'Video too large ({file_size} bytes, max {MAX_VIDEO_SIZE})'})

        # Create temp directory for frames and OCR output
        work_dir = tempfile.mkdtemp(prefix='video-ocr-')
        frames_dir = Path(work_dir) / 'frames'
        ocr_dir = Path(work_dir) / 'ocr'
        frames_dir.mkdir()
        ocr_dir.mkdir()

        try:
            # Extract frames using ffmpeg
            logger.info(f"Extracting frames at {fps} fps...")
            ffmpeg_cmd = [
                'ffmpeg', '-hide_banner', '-loglevel', 'warning',
                '-i', str(full_path),
                '-vf', f'fps={fps}',
                '-q:v', '2',
                str(frames_dir / 'frame_%06d.png')
            ]
            result = subprocess.run(ffmpeg_cmd, capture_output=True, text=True, timeout=300)
            if result.returncode != 0:
                logger.error(f"ffmpeg failed: {result.stderr}")
                return json.dumps({'success': False, 'error': f'Frame extraction failed: {result.stderr}'})

            frame_files = sorted(frames_dir.glob('frame_*.png'))
            logger.info(f"Extracted {len(frame_files)} frames")

            if not frame_files:
                return json.dumps({'success': False, 'error': 'No frames extracted from video'})

            # Run OCR on each frame
            logger.info(f"Running OCR on {len(frame_files)} frames...")
            for frame in frame_files:
                ocr_output = ocr_dir / frame.stem
                tesseract_cmd = ['tesseract', str(frame), str(ocr_output), '-l', 'eng', '--psm', '6']
                subprocess.run(tesseract_cmd, capture_output=True, timeout=60)

            # Read and deduplicate OCR results
            logger.info("Deduplicating text...")
            all_lines = []
            seen_lines = set()
            threshold = similarity_threshold / 100.0

            def similarity(a, b):
                return SequenceMatcher(None, a.lower(), b.lower()).ratio()

            ocr_files = sorted(ocr_dir.glob('*.txt'))
            for ocr_file in ocr_files:
                try:
                    with open(ocr_file, 'r', encoding='utf-8', errors='ignore') as f:
                        lines = [line.strip() for line in f if line.strip() and len(line.strip()) >= 3]

                    for line in lines:
                        normalized = line.lower().strip()
                        is_duplicate = normalized in seen_lines

                        if not is_duplicate:
                            # Check fuzzy match against recent lines
                            for seen in list(seen_lines)[-100:]:
                                if similarity(normalized, seen) > threshold:
                                    is_duplicate = True
                                    break

                        if not is_duplicate:
                            all_lines.append(line)
                            seen_lines.add(normalized)
                except Exception as e:
                    logger.debug(f"Error reading OCR file {ocr_file}: {e}")

            elapsed = time.time() - start_time
            logger.info(f"extract_video_text completed: {len(all_lines)} unique lines in {elapsed:.2f}s")

            return json.dumps({
                'success': True,
                'file_path': file_path,
                'frames_extracted': len(frame_files),
                'unique_lines': len(all_lines),
                'processing_time': f"{elapsed:.1f}s",
                'content': '\n'.join(all_lines)
            }, indent=2)

        finally:
            # Cleanup temp directory
            shutil.rmtree(work_dir, ignore_errors=True)

    except subprocess.TimeoutExpired:
        elapsed = time.time() - start_time
        logger.error(f"Video processing timed out after {elapsed:.2f}s")
        return json.dumps({'success': False, 'error': 'Video processing timed out'})
    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error extracting video text after {elapsed:.2f}s: {e}", exc_info=True)
        return json.dumps({'success': False, 'error': str(e)})


@mcp.tool()
async def transcribe_audio(
    file_path: str,
    language: str = "en",
    output_format: str = "txt"
) -> str:
    """
    Transcribe an audio file using the Whisper API. Sends the audio to a remote
    Whisper server and saves the transcript alongside the original file.

    Args:
        file_path: Relative path to the audio file
        language: Language code for transcription (default 'en')
        output_format: Output format - 'txt', 'vtt', 'srt', or 'json' (default 'txt')

    Returns:
        JSON string with transcription result and saved transcript path
    """
    start_time = time.time()
    logger.info(f"transcribe_audio called: file_path='{file_path}', language='{language}', format='{output_format}'")

    try:
        import json

        full_path = DOCUMENTS_PATH / file_path
        logger.debug(f"Full path resolved to: {full_path}")

        if not is_safe_path(DOCUMENTS_PATH, full_path):
            logger.warning(f"Path safety check failed for: {full_path}")
            return json.dumps({'success': False, 'error': 'Access denied - path outside allowed directory'})

        if not full_path.exists():
            logger.warning(f"File not found: {full_path}")
            return json.dumps({'success': False, 'error': f'File not found: {file_path}'})

        if full_path.suffix.lower() not in [e.lower() for e in AUDIO_EXTENSIONS]:
            return json.dumps({'success': False, 'error': f'Not an audio file: {full_path.suffix}'})

        file_size = full_path.stat().st_size
        rel_path = str(full_path.relative_to(DOCUMENTS_PATH))

        # Check if transcript already exists
        transcript_path = full_path.with_suffix(TRANSCRIPT_EXT)
        if transcript_path.exists():
            # Return existing transcript
            with open(transcript_path, 'r', encoding='utf-8', errors='ignore') as f:
                transcript_text = f.read()

            elapsed = time.time() - start_time
            logger.info(f"transcribe_audio returning existing transcript: {transcript_path.name} ({len(transcript_text)} chars) in {elapsed:.2f}s")

            return json.dumps({
                'success': True,
                'file_path': file_path,
                'transcript_path': str(transcript_path.relative_to(DOCUMENTS_PATH)),
                'transcript_size': len(transcript_text.encode('utf-8')),
                'audio_size': file_size,
                'status': 'completed',
                'content': transcript_text
            }, indent=2)

        # No transcript yet - check status
        is_in_progress = rel_path in _transcription_in_progress
        last_error = _transcription_errors.get(rel_path)

        if is_in_progress:
            elapsed = time.time() - start_time
            logger.info(f"transcribe_audio: transcription in progress for {rel_path}")
            return json.dumps({
                'success': True,
                'file_path': file_path,
                'audio_size': file_size,
                'status': 'in_progress',
                'message': 'Transcription is currently in progress. The background scanner is processing this file. Check back in a few minutes.'
            }, indent=2)

        # Not yet transcribed and not in progress
        elapsed = time.time() - start_time
        message = 'This audio file has not been transcribed yet. The background scanner will pick it up automatically.'
        if last_error:
            message += f' Last error: {last_error}'
        if not WHISPER_API_URL:
            message = 'Whisper API not configured. Set WHISPER_API_URL environment variable.'

        logger.info(f"transcribe_audio: no transcript for {rel_path}")
        return json.dumps({
            'success': False,
            'file_path': file_path,
            'audio_size': file_size,
            'status': 'pending',
            'message': message
        }, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error in transcribe_audio for {file_path} after {elapsed:.2f}s: {e}", exc_info=True)
        return json.dumps({'success': False, 'error': str(e)})


@mcp.tool()
async def download_file(file_path: str) -> str:
    """
    Get a direct download URL for a file. Works with any file type including videos.

    Args:
        file_path: Relative path to the file

    Returns:
        JSON string with download URL and file metadata
    """
    import urllib.parse

    start_time = time.time()
    logger.info(f"download_file called: file_path='{file_path}'")

    try:
        import json

        full_path = DOCUMENTS_PATH / file_path
        logger.debug(f"Full path resolved to: {full_path}")

        if not is_safe_path(DOCUMENTS_PATH, full_path):
            logger.warning(f"Path safety check failed for: {full_path}")
            return json.dumps({'success': False, 'error': 'Access denied - path outside allowed directory'})

        if not full_path.exists():
            logger.warning(f"File not found: {full_path}")
            return json.dumps({'success': False, 'error': f'File not found: {file_path}'})

        # Get file info
        file_size = full_path.stat().st_size
        logger.debug(f"File size: {file_size} bytes")

        # Determine MIME type
        mime_types = {
            '.mp4': 'video/mp4',
            '.mkv': 'video/x-matroska',
            '.avi': 'video/x-msvideo',
            '.webm': 'video/webm',
            '.mov': 'video/quicktime',
            '.m4v': 'video/x-m4v',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.json': 'application/json',
            '.csv': 'text/csv',
        }
        mime_type = mime_types.get(full_path.suffix.lower(), 'application/octet-stream')

        # Build download URL
        encoded_path = urllib.parse.quote(file_path, safe='')
        if PUBLIC_URL:
            download_url = f"{PUBLIC_URL}/download/{encoded_path}?token={AUTH_TOKEN}"
        else:
            download_url = f"/download/{encoded_path}?token={AUTH_TOKEN}"
            logger.warning("MCP_PUBLIC_URL not set - returning relative URL")

        elapsed = time.time() - start_time
        logger.info(f"download_file completed: {file_size} bytes, url generated in {elapsed:.2f}s")

        return json.dumps({
            'success': True,
            'file_path': file_path,
            'file_name': full_path.name,
            'file_size': file_size,
            'file_size_mb': round(file_size / 1024 / 1024, 1),
            'mime_type': mime_type,
            'download_url': download_url,
            'instructions': 'Open the download_url in a browser to download the file directly.',
            'note': 'If network is restricted, use download_file_chunk() instead to get the file as base64 data chunks.'
        }, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error downloading file {file_path} after {elapsed:.2f}s: {e}", exc_info=True)
        return json.dumps({'success': False, 'error': str(e)})


@mcp.tool()
async def download_file_chunk(
    file_path: str,
    chunk: int = 0,
    chunk_size_mb: int = 5
) -> str:
    """
    Download a file in chunks as base64-encoded data. Use this for large files.
    Call with chunk=0 first to get total_chunks, then iterate through all chunks.

    Args:
        file_path: Relative path to the file
        chunk: Chunk number (0-indexed)
        chunk_size_mb: Size of each chunk in MB (default 5MB)

    Returns:
        JSON string with base64-encoded chunk data and metadata
    """
    import base64
    import math

    start_time = time.time()
    chunk_size = chunk_size_mb * 1024 * 1024
    logger.info(f"download_file_chunk called: file_path='{file_path}', chunk={chunk}, chunk_size_mb={chunk_size_mb}")

    try:
        import json

        full_path = DOCUMENTS_PATH / file_path
        logger.debug(f"Full path resolved to: {full_path}")

        if not is_safe_path(DOCUMENTS_PATH, full_path):
            logger.warning(f"Path safety check failed for: {full_path}")
            return json.dumps({'success': False, 'error': 'Access denied - path outside allowed directory'})

        if not full_path.exists():
            logger.warning(f"File not found: {full_path}")
            return json.dumps({'success': False, 'error': f'File not found: {file_path}'})

        # Get file info
        file_size = full_path.stat().st_size
        total_chunks = math.ceil(file_size / chunk_size)

        logger.debug(f"File size: {file_size} bytes, total_chunks: {total_chunks}")

        # Validate chunk number
        if chunk < 0 or chunk >= total_chunks:
            return json.dumps({
                'success': False,
                'error': f'Invalid chunk {chunk}. File has {total_chunks} chunks (0-{total_chunks-1})',
                'total_chunks': total_chunks
            })

        # Calculate byte range for this chunk
        start_byte = chunk * chunk_size
        end_byte = min(start_byte + chunk_size, file_size)
        bytes_to_read = end_byte - start_byte

        logger.info(f"Reading chunk {chunk}/{total_chunks-1}: bytes {start_byte}-{end_byte} ({bytes_to_read} bytes)")

        # Read the chunk
        with open(full_path, 'rb') as f:
            f.seek(start_byte)
            chunk_data = f.read(bytes_to_read)

        # Encode to base64
        encoded_data = base64.b64encode(chunk_data).decode('ascii')

        elapsed = time.time() - start_time
        logger.info(f"download_file_chunk completed: chunk {chunk}/{total_chunks-1}, "
                   f"{len(chunk_data)} bytes -> {len(encoded_data)} chars in {elapsed:.2f}s")

        return json.dumps({
            'success': True,
            'file_path': file_path,
            'file_name': full_path.name,
            'file_size': file_size,
            'chunk': chunk,
            'total_chunks': total_chunks,
            'chunk_size_mb': chunk_size_mb,
            'start_byte': start_byte,
            'end_byte': end_byte,
            'bytes_in_chunk': len(chunk_data),
            'encoding': 'base64',
            'data': encoded_data,
            'is_last_chunk': chunk == total_chunks - 1
        }, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error downloading chunk {chunk} of {file_path} after {elapsed:.2f}s: {e}", exc_info=True)
        return json.dumps({'success': False, 'error': str(e)})


# Extensions that can be created/written
WRITABLE_EXTENSIONS = ['.txt', '.md', '.json', '.yaml', '.yml', '.csv', '.log']


@mcp.tool()
async def create_document(file_path: str, content: str, overwrite: bool = False) -> str:
    """
    Create a new document or update an existing one.

    Args:
        file_path: Relative path for the document (e.g., 'notes/meeting.txt')
        content: Text content to write to the file
        overwrite: If True, overwrite existing files. If False, fail if file exists.

    Returns:
        JSON string with result status and file info
    """
    start_time = time.time()
    content_size = len(content.encode('utf-8'))
    logger.info(f"create_document called: file_path='{file_path}', content_size={content_size} bytes, overwrite={overwrite}")

    try:
        import json

        # Validate file path
        if not file_path:
            logger.warning("Empty file path provided")
            return json.dumps({'success': False, 'error': 'File path cannot be empty'})

        full_path = DOCUMENTS_PATH / file_path
        logger.debug(f"Full path resolved to: {full_path}")

        # Security: ensure path is within documents directory
        if not is_safe_path(DOCUMENTS_PATH, full_path):
            logger.warning(f"Path safety check failed for: {full_path}")
            return json.dumps({'success': False, 'error': 'Invalid path - must be within documents directory'})

        # Check extension is writable
        if full_path.suffix not in WRITABLE_EXTENSIONS:
            logger.warning(f"Extension not writable: {full_path.suffix}")
            return json.dumps({
                'success': False,
                'error': f'Extension {full_path.suffix} not allowed. Allowed: {", ".join(WRITABLE_EXTENSIONS)}'
            })

        # Check if file exists
        file_existed = full_path.exists()
        if file_existed and not overwrite:
            logger.warning(f"File exists and overwrite=False: {full_path}")
            return json.dumps({
                'success': False,
                'error': f'File already exists: {file_path}. Set overwrite=True to replace it.'
            })

        # Create parent directories if needed
        if not full_path.parent.exists():
            logger.debug(f"Creating parent directories: {full_path.parent}")
            full_path.parent.mkdir(parents=True, exist_ok=True)

        # Write the file
        logger.debug(f"Writing {content_size} bytes to {full_path}")
        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(content)

        elapsed = time.time() - start_time
        logger.info(f"create_document completed: wrote {content_size} bytes to {file_path}, overwritten={file_existed and overwrite}, took {elapsed:.2f}s")

        return json.dumps({
            'success': True,
            'file_path': file_path,
            'size': content_size,
            'overwritten': file_existed and overwrite
        }, indent=2)

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"Error creating document {file_path} after {elapsed:.2f}s: {e}", exc_info=True)
        return json.dumps({'success': False, 'error': str(e)})


@mcp.resource("documents://list")
def list_documents_resource() -> str:
    """Resource that provides a list of all available documents"""
    documents = []
    for file_path in DOCUMENTS_PATH.rglob('*'):
        if file_path.is_file() and (file_path.suffix.lower() in [e.lower() for e in ALL_EXTENSIONS]):
            documents.append(get_file_info(file_path))
    
    import json
    return json.dumps({
        'total': len(documents),
        'documents': documents[:100]  # Limit to 100 for resource
    }, indent=2)


def main():
    """Main entry point"""
    import uvicorn
    from starlette.middleware.base import BaseHTTPMiddleware
    from starlette.responses import Response

    # Ensure documents directory exists
    DOCUMENTS_PATH.mkdir(parents=True, exist_ok=True)

    # Log startup configuration
    logger.info("=" * 60)
    logger.info("MCP Document Server Starting")
    logger.info("=" * 60)
    logger.info(f"Documents path: {DOCUMENTS_PATH}")
    logger.info(f"Documents path exists: {DOCUMENTS_PATH.exists()}")
    logger.info(f"Documents path writable: {os.access(DOCUMENTS_PATH, os.W_OK)}")
    logger.info(f"Max file size: {MAX_FILE_SIZE / 1024 / 1024:.1f} MB")
    logger.info(f"Allowed extensions: {ALLOWED_EXTENSIONS}")
    logger.info(f"Writable extensions: {WRITABLE_EXTENSIONS}")
    logger.info(f"Allowed hosts: {ALLOWED_HOSTS}")
    logger.info(f"Auth token configured: {bool(AUTH_TOKEN)}")
    logger.info(f"Whisper API URL: {WHISPER_API_URL or 'not configured'}")
    logger.info(f"Whisper scan interval: {WHISPER_SCAN_INTERVAL}s")
    logger.info(f"Transcript format: {TRANSCRIPT_EXT}")
    logger.info(f"Recordings directory: {RECORDINGS_DIR}")
    logger.info(f"Speakers directory: {RECORDINGS_DIR}/{SPEAKERS_DIR}")
    logger.info(f"Speaker similarity threshold: {SPEAKER_SIMILARITY_THRESHOLD}")
    logger.info(f"Log level: {LOG_LEVEL}")

    # Log speaker enrollment profiles
    try:
        profiles = _scan_speaker_profiles()
        if profiles:
            total_clips = sum(len(v) for v in profiles.values())
            logger.info(f"Found {len(profiles)} speaker profile(s) with {total_clips} total clip(s): {', '.join(sorted(profiles.keys()))}")
        else:
            logger.info("No speaker enrollment profiles found (speaker identification disabled)")
    except Exception as e:
        logger.warning(f"Could not scan speaker profiles: {e}")

    # Count existing documents
    try:
        doc_count = sum(1 for f in DOCUMENTS_PATH.rglob('*') if f.is_file() and f.suffix in ALLOWED_EXTENSIONS)
        logger.info(f"Found {doc_count} existing documents")
    except Exception as e:
        logger.warning(f"Could not count documents: {e}")

    # Check which transport to use based on environment
    transport = os.getenv('MCP_TRANSPORT', 'sse')

    if transport == 'stdio':
        logger.info("Using STDIO transport")
        mcp.run(transport='stdio')
    else:
        # Default to SSE for HTTP-based access
        host = os.getenv('MCP_HOST', '0.0.0.0')
        port = int(os.getenv('MCP_PORT', '8000'))
        logger.info(f"Using SSE transport on {host}:{port}")

        app = mcp.sse_app()

        # Add health endpoint
        from starlette.responses import JSONResponse, Response, FileResponse
        from starlette.routing import Route
        import urllib.parse

        # Start background audio scanner
        _scanner_task = None

        async def _start_scanner():
            nonlocal _scanner_task
            if WHISPER_API_URL:
                _scanner_task = asyncio.create_task(_audio_scanner_loop())
                logger.info("Background audio scanner task created")
            else:
                logger.info("Background audio scanner disabled (no WHISPER_API_URL)")

        app.add_event_handler("startup", _start_scanner)

        async def health(request):
            return JSONResponse({'status': 'healthy'})

        app.routes.append(Route('/health', health))

        # Add file download endpoint
        async def download(request):
            """Serve files directly for download"""
            # Get the file path from URL (everything after /download/)
            file_path_encoded = request.path_params.get('file_path', '')
            file_path = urllib.parse.unquote(file_path_encoded)

            logger.info(f"Download request: {file_path}")

            if not file_path:
                return Response(content="File path required", status_code=400)

            full_path = DOCUMENTS_PATH / file_path

            # Security check
            if not is_safe_path(DOCUMENTS_PATH, full_path):
                logger.warning(f"Download blocked - path outside allowed directory: {file_path}")
                return Response(content="Access denied", status_code=403)

            if not full_path.exists():
                logger.warning(f"Download failed - file not found: {file_path}")
                return Response(content="File not found", status_code=404)

            logger.info(f"Serving file: {full_path} ({full_path.stat().st_size} bytes)")
            return FileResponse(
                path=full_path,
                filename=full_path.name,
                media_type='application/octet-stream'
            )

        app.routes.append(Route('/download/{file_path:path}', download))

        # Add authentication middleware if token is configured
        if AUTH_TOKEN:
            logger.info(f"Authentication enabled (token length: {len(AUTH_TOKEN)})")

            # Use pure ASGI middleware (BaseHTTPMiddleware breaks SSE streaming)
            from starlette.datastructures import URL, QueryParams

            class AuthMiddleware:
                def __init__(self, app):
                    self.app = app

                async def __call__(self, scope, receive, send):
                    if scope["type"] != "http":
                        return await self.app(scope, receive, send)

                    path = scope.get("path", "")
                    method = scope.get("method", "UNKNOWN")
                    client = scope.get("client", ("unknown", 0))
                    client_ip = client[0] if client else "unknown"

                    logger.debug(f"Request: {method} {path} from {client_ip}")

                    # Allow healthcheck and messages endpoints without token
                    # (messages require valid session_id from authenticated SSE)
                    if path == "/health":
                        logger.debug(f"Health check from {client_ip}")
                        return await self.app(scope, receive, send)

                    if path.startswith("/messages"):
                        logger.debug(f"Messages endpoint from {client_ip} (session auth)")
                        return await self.app(scope, receive, send)

                    # Check Authorization header
                    headers = dict(scope.get("headers", []))
                    auth_header = headers.get(b"authorization", b"").decode()
                    if auth_header == f"Bearer {AUTH_TOKEN}":
                        logger.debug(f"Authenticated via Bearer token from {client_ip}")
                        return await self.app(scope, receive, send)

                    # Check token query parameter
                    query_string = scope.get("query_string", b"").decode()
                    query_params = QueryParams(query_string)
                    if query_params.get("token") == AUTH_TOKEN:
                        logger.debug(f"Authenticated via query param from {client_ip}")
                        return await self.app(scope, receive, send)

                    # Unauthorized - log details for debugging
                    has_auth_header = bool(auth_header)
                    has_query_token = "token" in query_params
                    logger.warning(
                        f"Unauthorized request: {method} {path} from {client_ip} "
                        f"(has_auth_header={has_auth_header}, has_query_token={has_query_token})"
                    )
                    response = Response(content="Unauthorized", status_code=401)
                    await response(scope, receive, send)

            app = AuthMiddleware(app)
        else:
            logger.warning("No MCP_AUTH_TOKEN set - server is unauthenticated!")

        uvicorn.run(app, host=host, port=port)


if __name__ == '__main__':
    main()
